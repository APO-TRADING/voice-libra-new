"""
Text extraction and cleaning utilities for ebook files.
Ported from user-provided text-converter-cleaner-v5.py.
Supports PDF, EPUB, DOCX, TXT.
"""
import os
import re
import warnings
from io import BytesIO

warnings.filterwarnings("ignore", category=UserWarning, module='bs4')


class DocumentProcessor:
    def extract(self, filename: str, content: bytes) -> str:
        ext = os.path.splitext(filename)[1].lower()
        if ext == '.pdf':
            return self._from_pdf(content)
        elif ext == '.epub':
            return self._from_epub(content)
        elif ext == '.docx':
            return self._from_docx(content)
        elif ext == '.txt':
            return self._from_txt(content)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    def _from_pdf(self, content: bytes) -> str:
        text = ""
        try:
            import pdfplumber
            with pdfplumber.open(BytesIO(content)) as pdf:
                for page in pdf.pages:
                    text += (page.extract_text() or "") + "\n\n"
        except Exception:
            import PyPDF2
            reader = PyPDF2.PdfReader(BytesIO(content))
            for page in reader.pages:
                text += (page.extract_text() or "") + "\n\n"
        return text.strip()

    def _from_epub(self, content: bytes) -> str:
        # ebooklib needs a file path, so write temp file
        import tempfile
        from ebooklib import epub
        from bs4 import BeautifulSoup
        with tempfile.NamedTemporaryFile(suffix=".epub", delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        try:
            book = epub.read_epub(tmp_path)
            text = ""
            for item in book.get_items():
                if item.get_type() == 9:
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    txt = soup.get_text(separator=' ')
                    text += re.sub(r'\s+', ' ', txt).strip() + "\n\n"
            return text.strip()
        finally:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass

    def _from_docx(self, content: bytes) -> str:
        import docx
        doc = docx.Document(BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    def _from_txt(self, content: bytes) -> str:
        for enc in ('utf-8', 'latin-1', 'cp1252'):
            try:
                return content.decode(enc)
            except UnicodeDecodeError:
                continue
        return content.decode('utf-8', errors='replace')


class TextCleaner:
    def clean(self, text: str) -> str:
        text = re.sub(r'-\s*\n', '', text)
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n{2,}', '\n\n', text)
        text = self._join_broken(text)
        text = self._fix_punct(text)
        text = self._special(text)
        text = re.sub(r' +', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = self._add_paragraph_break(text)
        text = self._join_paragraphs(text)
        text = self._remove_page_numbers(text)
        return text.strip()

    def _add_paragraph_break(self, text: str) -> str:
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(
            r'([.!?])([»”"]?)\s+([A-ZÀ-ÈÌÒÙ])',
            r'\1\2\n\n\3', text
        )
        text = re.sub(
            r'\.{3}([»”"]?)\s+([A-ZÀ-ÈÌÒÙ])',
            r'...\1\n\n\2', text
        )
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text

    def _join_broken(self, text: str) -> str:
        lines = text.splitlines()
        result = []
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                result.append('')
                i += 1
                continue
            if i == len(lines) - 1 or re.search(r'[.!?:"»]$', line):
                result.append(line)
            else:
                next_line = lines[i + 1].strip()
                if next_line and (next_line[0].islower() or next_line[0] in ',:;)]}'):
                    result.append(f"{line} {next_line}")
                    i += 1
                else:
                    result.append(line)
            i += 1
        return '\n'.join(result)

    def _fix_punct(self, text: str) -> str:
        text = re.sub(r'\.([»”"])', r'\1.', text)
        text = re.sub(r'([.!?:;,])([^\s»""])(?!\d)', r'\1 \2', text)
        text = re.sub(r'(\s)([.!?:;,])', r'\2', text)
        text = re.sub(r'([«""])\s+', r'\1', text)
        text = re.sub(r'\s+([»""])', r'\1', text)
        text = re.sub(r'\.\s*\.\s*\.', '...', text)
        return text

    def _special(self, text: str) -> str:
        text = re.sub(r'[—–]', '-', text)
        text = re.sub(r'^[\s•\-*]+', '', text, flags=re.MULTILINE)
        text = re.sub(r'-{2,}', '-', text)
        return text

    def _join_paragraphs(self, text: str, min_length: int = 40) -> str:
        paragraphs = text.split('\n\n')
        result = []
        i = 0
        while i < len(paragraphs):
            current = paragraphs[i].strip()
            if not current:
                i += 1
                continue
            if len(current) < min_length and i < len(paragraphs) - 1:
                is_heading = current.endswith(':') or current.isupper()
                if not is_heading:
                    nxt = paragraphs[i + 1].strip()
                    if nxt:
                        result.append(f"{current} {nxt}")
                        i += 2
                        continue
            result.append(current)
            i += 1
        return '\n\n'.join(result)

    def _remove_page_numbers(self, text: str) -> str:
        text = re.sub(r'^\s*\d+\s*$', '', text, flags=re.MULTILINE)
        text = re.sub(r'^\s*Page \d+ of \d+\s*$', '', text, flags=re.MULTILINE)
        text = re.sub(r'^\s*\d+\s*\|\s*Page\s*$', '', text, flags=re.MULTILINE)
        return re.sub(r'\n{3,}', '\n\n', text)


# Sentence splitter compatible with the cleaner output (Italian-aware)
SENTENCE_RE = re.compile(
    r'[^.!?…]+(?:[.!?…]+["»”\')\]]*|$)',
    re.DOTALL
)


def split_sentences(text: str) -> list[str]:
    """Split cleaned text into sentences for sentence-level navigation."""
    sentences: list[str] = []
    for raw in SENTENCE_RE.findall(text):
        s = raw.strip()
        if s:
            sentences.append(s)
    return sentences


def count_words(text: str) -> int:
    return len(re.findall(r'\b\w+\b', text))
